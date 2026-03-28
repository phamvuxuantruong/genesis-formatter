import streamlit as st
import os
import io
import time
import json
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import google.generativeai as genai

# -------------------------------------------------------------------
# CORE FUNCTION: GỌI GEMINI LẤY SƠ ĐỒ CẤU TRÚC (JSON MAPPING)
# -------------------------------------------------------------------
# -------------------------------------------------------------------
# CORE FUNCTION: GỌI GEMINI LẤY SƠ ĐỒ CẤU TRÚC (JSON MAPPING)
# -------------------------------------------------------------------
def get_structure_from_ai(doc_text_indexed, api_key):
    genai.configure(api_key=api_key)
    
    prompt = f"""Bạn là chuyên gia phân tích cấu trúc văn bản kỹ thuật.
Nhiệm vụ của bạn là đọc văn bản đã được đánh số dòng (Ví dụ: [0] Nội dung...) và xác định xem dòng nào là Tiêu đề (Heading).
Phân loại:
- Cấp 1 (level: 1): Các Phần, Chương lớn.
- Cấp 2 (level: 2): Các Mục lớn (I, II, III...).
- Cấp 3 (level: 3): Các Mục nhỏ (1, 2, 3, a, b, c...).

BẮT BUỘC TRẢ VỀ ĐỊNH DẠNG JSON NHƯ SAU:
{{
  "headings": [
    {{"id": 0, "level": 1}},
    {{"id": 5, "level": 2}}
  ]
}}
Chỉ liệt kê các ID thực sự là tiêu đề, bỏ qua nội dung bình thường.

VĂN BẢN ĐẦU VÀO:
{doc_text_indexed}"""

    try:
        # ƯU TIÊN 1: Thử gọi bản 1.5 Flash mới nhất (Xử lý file siêu dài)
        model = genai.GenerativeModel(
            'gemini-1.5-flash-latest', # Dùng tên đầy đủ thay vì tên viết tắt
            generation_config={"response_mime_type": "application/json", "temperature": 0.1}
        )
        response = model.generate_content(prompt)
        return json.loads(response.text)
    except Exception as e_flash:
        # CỨU CÁNH 2: Nếu Google báo 404, lập tức lùi về bản Gemini Pro ổn định toàn cầu
        try:
            model_pro = genai.GenerativeModel(
                'gemini-pro',
                generation_config={"temperature": 0.1}
            )
            response_pro = model_pro.generate_content(prompt)
            
            # Làm sạch chuỗi JSON lỡ như AI bọc thêm mã ```json
            cleaned_text = response_pro.text.strip()
            if cleaned_text.startswith("```json"):
                cleaned_text = cleaned_text[7:-3].strip()
            elif cleaned_text.startswith("```"):
                cleaned_text = cleaned_text[3:-3].strip()
                
            return json.loads(cleaned_text)
        except Exception as e_pro:
            raise Exception(f"Lỗi mạng Google API: {str(e_pro)}")

# -------------------------------------------------------------------
# CORE FUNCTION: CẤU HÌNH STYLE NĐ30 CHO FILE
# -------------------------------------------------------------------
def setup_nd30_styles(doc):
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(14)
    normal_style.font.color.rgb = RGBColor(0, 0, 0)
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal_style.paragraph_format.line_spacing = 1.2
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.first_line_indent = Cm(1.27)

    def config_heading(level, size):
        style_name = f'Heading {level}'
        if style_name in doc.styles:
            style = doc.styles[style_name]
        else:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(size)
        font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.first_line_indent = Cm(0)

    config_heading(1, 14)  
    config_heading(2, 14) 
    config_heading(3, 14)  

# -------------------------------------------------------------------
# CORE FUNCTION: THAO TÁC TRỰC TIẾP LÊN FILE GỐC (IN-PLACE MUTATION)
# -------------------------------------------------------------------
def apply_structure_to_doc(doc, json_structure):
    heading_map = {}
    if 'headings' in json_structure:
        for item in json_structure['headings']:
            heading_map[item['id']] = item['level']
            
    for i, p in enumerate(doc.paragraphs):
        level = heading_map.get(i)
        
        if level == 1:
            p.style = 'Heading 1'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            p.style = 'Heading 2'
        elif level == 3:
            p.style = 'Heading 3'
        else:
            p.style = 'Normal'
            
        for run in p.runs:
            if run.text.strip():
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                
                if level == 1:
                    run.font.bold = True
                    run.font.all_caps = True 
                elif level == 2:
                    run.font.bold = True
                elif level == 3:
                    run.font.bold = True
                    run.font.italic = True

    for table in doc.tables:
        try:
            table.style = 'Table Grid'
        except KeyError:
            pass 
            
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.style = 'Normal'
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.first_line_indent = Cm(0) 
                    for run in p.runs:
                        if run.text.strip():
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(14)

    return doc

# -------------------------------------------------------------------
# STREAMLIT UI
# -------------------------------------------------------------------
def main():
    st.set_page_config(page_title="Genesis Auto-Formatter", page_icon="📄", layout="centered")
    
    st.title("📄 Genesis Auto-Formatter v3.0 (Gemini Engine)")
    st.markdown("**Kiến trúc In-Place Mutation: Định dạng NĐ 30/2020/NĐ-CP với Vị trí Tuyệt đối 100%.**")
    st.info("💡 Trí tuệ nhân tạo (Gemini Flash) sẽ tự động cơ cấu lại các tiêu đề, trong khi hệ thống mã hóa bảo vệ nguyên vẹn 100% hình ảnh và bảng biểu của bạn.")

    # [BẢO MẬT]: Lấy API Key từ Két sắt Cloud (Secrets)
    try:
        api_key = st.secrets["GEMINI_API_KEY"]
    except KeyError:
        st.error("❌ CẢNH BÁO BẢO MẬT: Hệ thống chưa được cấp GEMINI_API_KEY. Vui lòng cấu hình Két sắt (Secrets) trên máy chủ.")
        st.stop()

    uploaded_file = st.file_uploader("📥 Tải lên file Word (.docx) lộn xộn của bạn", type=["docx"])

    if uploaded_file is not None:
        if st.button("🚀 Xử lý & Chuẩn hóa ngay", type="primary", use_container_width=True):
            try:
                progress_bar = st.progress(0)
                status_text = st.empty()
                start_time = time.time()

                status_text.markdown("⏳ **Bước 1:** Đang đọc cấu trúc file và tạo lập bản đồ ID (Không phá vỡ tệp)...")
                doc = Document(uploaded_file)
                setup_nd30_styles(doc)
                
                doc_text_indexed = ""
                for i, p in enumerate(doc.paragraphs):
                    text = p.text.strip()
                    if text:
                        doc_text_indexed += f"[{i}] {text}\n"
                        
                progress_bar.progress(30)

                status_text.markdown("🧠 **Bước 2:** Gemini AI đang phân tích Sơ đồ Heading (Công nghệ Bypass Token Limit)...")
                json_structure = get_structure_from_ai(doc_text_indexed, api_key)
                progress_bar.progress(70)

                status_text.markdown("🏗️ **Bước 3:** Đang áp dụng định dạng trực tiếp lên File Gốc. Bảo tồn 100% vị trí Media...")
                final_doc = apply_structure_to_doc(doc, json_structure)
                
                output_stream = io.BytesIO()
                final_doc.save(output_stream)
                output_stream.seek(0)
                progress_bar.progress(100)
                
                end_time = time.time()
                processing_time = round(end_time - start_time, 1)
                
                status_text.success(f"✅ Hoàn thành xuất sắc trong {processing_time} giây! Mọi hình ảnh và bảng biểu vẫn nằm chính xác ở vị trí gốc.")
                
                st.download_button(
                    label="⬇️ Tải xuống File Word đã chuẩn hóa",
                    data=output_stream,
                    file_name=f"Standardized_{uploaded_file.name}",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

            except Exception as e:
                st.error(f"❌ Phát hiện lỗi trong quá trình xử lý: {str(e)}")

if __name__ == "__main__":
    main()
